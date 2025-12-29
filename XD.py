import tkinter as tk
from tkinter import messagebox, filedialog
import random
from docx import Document
import os
import urllib.request

# آدرس مستقیم فایل لیست کلمات
BIP39_URL = "https://raw.githubusercontent.com/bitcoin/bips/master/bip-0039/english.txt"
BIP39_FILE = "bip39_english.txt"

# تابع دانلود فایل
def download_bip39_file():
    try:
        urllib.request.urlretrieve(BIP39_URL, BIP39_FILE)
        print("✅ فایل bip39_english.txt با موفقیت دانلود شد.")
    except Exception as e:
        messagebox.showerror("خطا", f"دانلود فایل ناموفق بود:\n{e}")

# تابع بارگذاری فایل
def load_bip39_words():
    if not os.path.isfile(BIP39_FILE):
        print("⚠️ فایل پیدا نشد، در حال دانلود...")
        download_bip39_file()

    if not os.path.isfile(BIP39_FILE):
        messagebox.showerror("خطا", "فایل bip39_english.txt وجود ندارد و دانلود هم نشد.")
        return []

    with open(BIP39_FILE, "r", encoding="utf-8") as file:
        words = [line.strip() for line in file.readlines() if line.strip()]
        if len(words) != 2048:
            messagebox.showerror("خطا", "تعداد کلمات فایل نادرست است. باید دقیقاً 2048 کلمه باشد.")
            return []
        return words

# بارگذاری کلمات
BIP39_WORDS = load_bip39_words()
generated_combinations = []

def generate_random_combinations(words, combo_size=12, count=1000):
    combinations = set()
    attempts = 0
    max_attempts = count * 10

    while len(combinations) < count and attempts < max_attempts:
        combo = tuple(random.sample(words, combo_size))
        combinations.add(combo)
        attempts += 1

    if len(combinations) < count:
        messagebox.showwarning("هشدار", "تعداد ترکیب‌های تولید شده کمتر از حد مورد نظر است.")
    return list(combinations)

def on_generate_button_click():
    global generated_combinations

    if not BIP39_WORDS:
        return

    generated_combinations = generate_random_combinations(BIP39_WORDS, combo_size=12, count=1000)
    result_text.delete(1.0, tk.END)
    for combo in generated_combinations[:30]:
        result_text.insert(tk.END, " ".join(combo) + "\n\n")

def save_as_word():
    if not generated_combinations:
        messagebox.showerror("خطا", "ابتدا ترکیب‌ها را تولید کنید.")
        return

    file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                             filetypes=[("Word files", "*.docx")],
                                             title="ذخیره Word")
    if not file_path:
        return

    doc = Document()
    doc.add_heading('ترکیب‌های تصادفی ۱۲تایی (BIP-39)', 0)
    for i, combo in enumerate(generated_combinations, 1):
        doc.add_paragraph(f"{i}. " + " ".join(combo))
    doc.save(file_path)
    messagebox.showinfo("موفقیت", "فایل Word با موفقیت ذخیره شد.")

# رابط گرافیکی
root = tk.Tk()
root.title("تولید عبارت‌های بازیابی تصادفی BIP-39")
root.geometry("650x550")

generate_button = tk.Button(root, text="تولید ترکیب‌ها", command=on_generate_button_click)
generate_button.pack(pady=10)

save_word_button = tk.Button(root, text="ذخیره به صورت Word", command=save_as_word)
save_word_button.pack(pady=5)

result_text = tk.Text(root, height=20, width=80)
result_text.pack(pady=10)

root.mainloop()
